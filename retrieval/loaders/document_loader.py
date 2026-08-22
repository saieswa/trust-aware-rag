"""
Document Loader and Structured Extractors.

Extracts text from PDF, DOCX, TXT, CSV, JSON, XLSX, and URLs while detecting
structural sections (Abstract, Introduction, Problem Statement, Methodology,
Experiments, Results, Conclusion, References, Appendix) and filtering out
repetitive header/footer noise.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse

from loguru import logger

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".doc", ".docx", ".csv", ".json", ".xlsx"}

# Academic & technical section patterns (supporting both standard and LaTeX-spaced headers)
SECTION_PATTERNS = [
    (re.compile(r"^(?:A\s*B\s*S\s*T\s*R\s*A\s*C\s*T)\b", re.I), "Abstract"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:I\s*N\s*T\s*R\s*O\s*D\s*U\s*C\s*T\s*I\s*O\s*N|B\s*A\s*C\s*K\s*G\s*R\s*O\s*U\s*N\s*D|M\s*O\s*T\s*I\s*V\s*A\s*T\s*I\s*O\s*N)\b", re.I), "Introduction"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:problem\s+(?:statement|formulation|definition)|research\s+gap)\b", re.I), "Problem Statement"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:R\s*E\s*L\s*A\s*T\s*E\s*D\s+W\s*O\s*R\s*K|P\s*R\s*I\s*O\s*R\s+W\s*O\s*R\s*K|L\s*I\s*T\s*E\s*R\s*A\s*T\s*U\s*R\s*E)\b", re.I), "Related Work"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:M\s*E\s*T\s*H\s*O\s*D\s*S?|M\s*E\s*T\s*H\s*O\s*D\s*O\s*L\s*O\s*G\s*Y|A\s*P\s*P\s*R\s*O\s*A\s*C\s*H|F\s*R\s*A\s*M\s*E\s*W\s*O\s*R\s*K|A\s*R\s*C\s*H\s*I\s*T\s*E\s*C\s*T\s*U\s*R\s*E|E\s*M\s*P\s*I\s*R\s*I\s*C\s*A\s*L\s+S\s*T\s*U\s*D\s*Y)\b", re.I), "Methodology"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:E\s*X\s*P\s*E\s*R\s*I\s*M\s*E\s*N\s*T\s*S?|E\s*V\s*A\s*L\s*U\s*A\s*T\s*I\s*O\s*N|S\s*E\s*T\s*U\s*P)\b", re.I), "Experiments"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:R\s*E\s*S\s*U\s*L\s*T\s*S?|F\s*I\s*N\s*D\s*I\s*N\s*G\s*S?|P\s*E\s*R\s*F\s*O\s*R\s*M\s*A\s*N\s*C\s*E|D\s*I\s*S\s*C\s*U\s*S\s*S\s*I\s*O\s*N)\b", re.I), "Results & Discussion"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:L\s*I\s*M\s*I\s*T\s*A\s*T\s*I\s*O\s*N\s*S?|F\s*U\s*T\s*U\s*R\s*E\s+W\s*O\s*R\s*K|B\s*R\s*O\s*A\s*D\s*E\s*R\s+I\s*M\s*P\s*A\s*C\s*T\s*S?)\b", re.I), "Limitations"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:C\s*O\s*N\s*C\s*L\s*U\s*S\s*I\s*O\s*N\s*S?|S\s*U\s*M\s*M\s*A\s*R\s*Y)\b", re.I), "Conclusion"),
    (re.compile(r"^(?:\d+[\.\s]+)?(?:R\s*E\s*F\s*E\s*R\s*E\s*N\s*C\s*E\s*S|B\s*I\s*B\s*L\s*I\s*O\s*G\s*R\s*A\s*P\s*H\s*Y)\b", re.I), "References"),
    (re.compile(r"^(?:A\s*P\s*P\s*E\s*N\s*D\s*I\s*X|[A-Z]\s+(?:D\s*E\s*T\s*A\s*I\s*L\s*S?|P\s*R\s*O\s*M\s*P\s*T|C\s*A\s*S\s*E\s+S\s*T\s*U\s*D\s*Y|P\s*R\s*O\s*O\s*F\s*S?))\b", re.I), "Appendix"),
]

HEADER_FOOTER_PATTERNS = [
    re.compile(r"Published as a conference paper at .*", re.I),
    re.compile(r"Under review as a conference paper at .*", re.I),
    re.compile(r"Preprint\. Under review\.*", re.I),
    re.compile(r"arXiv:\d+\.\d+(?:v\d+)?\s+\[[a-zA-Z\.\-]+\]\s+\d+\s+[a-zA-Z]+\s+\d{4}", re.I),
]


@dataclass
class Document:
    doc_id: str
    source_path: str
    title: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: List[Dict[str, Any]] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.content)


def _make_doc_id(identifier: str) -> str:
    digest = hashlib.sha256(identifier.strip().encode("utf-8")).hexdigest()
    return f"doc_{digest[:12]}"


def _normalize_latex_spacing(text: str) -> str:
    """Fix common PDF/LaTeX artifacts like 'D ETECTING' -> 'DETECTING' or 'M ETHODS' -> 'METHODS'."""
    text = re.sub(r"\b([A-Z])\s+([A-Z]{2,})\b", r"\1\2", text)
    text = re.sub(r"\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b", r"\1\2\3\4", text)
    return text


def _clean_page_text(text: str) -> str:
    """Strip repeated running headers/footers and normalize spaces."""
    if not text:
        return ""
    for pattern in HEADER_FOOTER_PATTERNS:
        text = pattern.sub("", text)

    text = _normalize_latex_spacing(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    lines = []
    for line in text.splitlines():
        clean_l = line.strip()
        if clean_l.isdigit():
            continue
        lines.append(clean_l)

    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def detect_section_header(line: str) -> Optional[str]:
    """Checks if a line acts as a section header."""
    clean = line.strip()
    if len(clean) > 80 or len(clean) < 3:
        return None
    for pattern, section_name in SECTION_PATTERNS:
        if pattern.search(clean):
            return section_name
    return None


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, List[Dict[str, Any]]]:
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    full_text_parts: List[str] = []
    pages_data: List[Dict[str, Any]] = []
    current_section = "General / Abstract"

    for page_idx, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
            cleaned = _clean_page_text(page_text)
            if not cleaned:
                continue

            # Detect sections on this page
            page_sections: List[str] = []
            for line in cleaned.splitlines():
                detected = detect_section_header(line)
                if detected:
                    current_section = detected
                    page_sections.append(detected)

            full_text_parts.append(f"[Page {page_idx} | Section: {current_section}]\n{cleaned}")
            pages_data.append({
                "page_number": page_idx,
                "section": current_section,
                "sections_found": page_sections,
                "text": cleaned,
            })
        except Exception as exc:
            logger.warning(f"Error extracting page {page_idx} from PDF: {exc}")

    full_text = "\n\n".join(full_text_parts)
    return full_text, pages_data


def extract_text_from_docx(file_bytes: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs: List[str] = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            paragraphs.append(txt)

    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                paragraphs.append(" | ".join(row_cells))

    return "\n\n".join(paragraphs)


def extract_text_from_csv(file_bytes: bytes) -> str:
    text_content = file_bytes.decode("utf-8", errors="replace")
    f = io.StringIO(text_content)
    reader = csv.DictReader(f)
    lines: List[str] = []

    for row_idx, row in enumerate(reader, start=1):
        items = [f"{k.strip()}: {v.strip()}" for k, v in row.items() if k and v is not None and v.strip()]
        if items:
            lines.append(f"Row {row_idx}: " + ", ".join(items) + ".")

    if not lines and text_content.strip():
        lines = [line.strip() for line in text_content.splitlines() if line.strip()]

    return "\n".join(lines)


def extract_text_from_json(file_bytes: bytes) -> str:
    raw_str = file_bytes.decode("utf-8", errors="replace")
    try:
        data = json.loads(raw_str)
    except Exception:
        return raw_str.strip()

    lines: List[str] = []
    if isinstance(data, list):
        for idx, item in enumerate(data, start=1):
            if isinstance(item, dict):
                entries = [f"{k}: {v}" for k, v in item.items() if v is not None]
                lines.append(f"Record {idx}: " + "; ".join(entries) + ".")
            else:
                lines.append(f"Item {idx}: {item}")
    elif isinstance(data, dict):
        for k, v in data.items():
            if isinstance(v, (dict, list)):
                lines.append(f"{k}: {json.dumps(v)}")
            else:
                lines.append(f"{k}: {v}")
    else:
        lines.append(str(data))

    return "\n".join(lines)


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines: List[str] = []

    for sheetname in wb.sheetnames:
        sheet = wb[sheetname]
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(cell).strip() if cell is not None else f"Col{idx}" for idx, cell in enumerate(rows[0])]

        for row_idx, row in enumerate(rows[1:], start=1):
            items = []
            for col_name, val in zip(headers, row):
                if val is not None and str(val).strip():
                    items.append(f"{col_name}: {str(val).strip()}")
            if items:
                lines.append(f"[{sheetname} - Row {row_idx}] " + ", ".join(items) + ".")

    return "\n".join(lines)


def extract_text_from_url(url: str, timeout: float = 15.0) -> tuple[str, str]:
    import httpx
    from bs4 import BeautifulSoup

    parsed = urlparse(url)
    if not parsed.scheme or parsed.scheme not in ("http", "https"):
        raise ValueError(f"Invalid URL scheme '{parsed.scheme}'. Must be http or https.")
    if not parsed.netloc:
        raise ValueError("Invalid URL: missing host domain.")

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) TrustAwareRAG/1.0",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }

    with httpx.Client(timeout=timeout, follow_redirects=True) as client:
        response = client.get(url, headers=headers)
        response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    page_title = soup.title.string.strip() if (soup.title and soup.title.string) else ""

    for tag in soup(["script", "style", "nav", "header", "footer", "aside", "svg", "noscript", "iframe"]):
        tag.decompose()

    main_el = soup.find("main") or soup.find("article") or soup.find("body") or soup
    cleaned_text = _clean_page_text(main_el.get_text(separator="\n"))

    if not cleaned_text:
        raise ValueError("Could not extract readable text from webpage.")

    if not page_title:
        first_line = next((line.strip() for line in cleaned_text.splitlines() if line.strip()), "")
        page_title = first_line[:80] if first_line else url

    return cleaned_text, page_title


def _derive_title(filename_or_url: str, content: str) -> str:
    first_line = next((line.strip() for line in content.splitlines() if line.strip()), "")
    if first_line and len(first_line) < 120 and not first_line.startswith("http") and not first_line.startswith("[Page"):
        clean_first = re.sub(r"^#+\s*", "", first_line).strip()
        if clean_first:
            return clean_first

    if filename_or_url.startswith("http://") or filename_or_url.startswith("https://"):
        parsed = urlparse(filename_or_url)
        return parsed.path.strip("/") or parsed.netloc

    stem = Path(filename_or_url).stem
    return stem.replace("_", " ").replace("-", " ").title()


def load_document_from_bytes(
    content_bytes: bytes,
    filename: str,
    source_url: Optional[str] = None,
) -> Document:
    if not content_bytes:
        raise ValueError(f"File '{filename}' is empty.")

    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    pages_data: List[Dict[str, Any]] = []
    file_type = ext.lstrip(".")

    if ext == ".pdf":
        text, pages_data = extract_text_from_pdf(content_bytes)
        file_type = "pdf"
    elif ext in (".doc", ".docx"):
        text = extract_text_from_docx(content_bytes)
        file_type = "docx"
    elif ext == ".csv":
        text = extract_text_from_csv(content_bytes)
        file_type = "csv"
    elif ext == ".json":
        text = extract_text_from_json(content_bytes)
        file_type = "json"
    elif ext == ".xlsx":
        text = extract_text_from_xlsx(content_bytes)
        file_type = "xlsx"
    else:
        text = _clean_page_text(content_bytes.decode("utf-8", errors="replace"))
        file_type = "txt"

    if not text.strip():
        raise ValueError(f"No extractable text found in '{filename}'.")

    doc_id = _make_doc_id(source_url or filename)
    title = _derive_title(filename, text)

    return Document(
        doc_id=doc_id,
        source_path=source_url or filename,
        title=title,
        content=text,
        metadata={
            "filename": filename,
            "source_type": "url" if source_url else "file",
            "source_url": source_url,
            "file_type": file_type,
            "size_bytes": len(content_bytes),
        },
        pages=pages_data,
    )


def load_document_from_url(url: str) -> Document:
    text, title = extract_text_from_url(url)
    doc_id = _make_doc_id(url)

    return Document(
        doc_id=doc_id,
        source_path=url,
        title=title,
        content=text,
        metadata={
            "filename": url,
            "source_type": "url",
            "source_url": url,
            "file_type": "url",
            "size_bytes": len(text.encode("utf-8")),
        },
        pages=[],
    )


def load_document(path: Path) -> Document:
    return load_document_from_bytes(path.read_bytes(), path.name)


def load_documents_from_directory(directory: str | Path) -> List[Document]:
    directory = Path(directory)
    if not directory.exists():
        logger.warning(f"Document directory does not exist: {directory}")
        return []

    documents: List[Document] = []
    for path in sorted(directory.iterdir()):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                documents.append(load_document(path))
            except Exception as exc:
                logger.error(f"Failed to load document {path}: {exc}")

    logger.info(f"Loaded {len(documents)} document(s) from {directory}")
    return documents
